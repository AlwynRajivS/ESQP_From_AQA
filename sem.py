import io, re, random, streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from copy import deepcopy
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from io import BytesIO
from docx.shared import Pt

# ==================================================
# STREAMLIT SETUP
# ==================================================
st.set_page_config(layout="wide")
st.title("KCET – Question Paper Generator")

qb_file = st.file_uploader("Upload AQA Question Bank (DOCX)", ["docx"])
template_file = st.file_uploader("Upload KCET Question Paper Template (DOCX)", ["docx"])

if not qb_file or not template_file:
    st.stop()

st.header("Exam Details")

month_duration = st.text_input(
    "Exam Duration (Month & Year)",
    placeholder="e.g. APR’2026"
)

semester = st.selectbox(
    "Semester",
    ["First", "Second", "Third", "Fourth", "Fifth", "Sixth", "Seventh"]
)

qp_code = st.text_input( "Question Paper Code", placeholder="e.g. VEC311-A" )

BOLD_PHRASES = [
    "B.E. / B.TECH. DEGREE EXAMINATIONS",
    "DEPARTMENT OF",
    "Reg. No.",
    "Roll No.",
    "Semester:",
    "Regulation:",
    "Max. Marks",
    "Duration",
    "CO Index",
    "Course Outcomes",
    "Marks distribution based on Bloom’s Taxonomy Level",
    "K-1",
    "K-2",
    "K-3",
    "K-4",
    "Total",
    "Answer all the Questions",
    "CO, BTL",
    "Part A",
    "Part B",
    "Part C",
    "Marks",
    "Q. No."
]

def clone_cell_xml(src_xml, dst_cell):
    dst_cell._element.clear_content()
    for el in src_xml:
        dst_cell._element.append(deepcopy(el))

def extract_aqa_metadata(file):
    doc = Document(file)

    all_text = []

    for p in doc.paragraphs:
        all_text.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)

    text = " ".join(all_text)
    text = re.sub(r"\s+", " ", text)  # normalize spaces

    data = {
        "department": "",
        "course_code": "",
        "course_name": ""
    }

    # Department: stop at "Course Code" or "Semester"
    d = re.search(
        r"Department:\s*(.+?)(?=\s+Course Code:|\s+Semester:)",
        text,
        re.I
    )
    if d:
        data["department"] = d.group(1).strip()

    # Course Code: stop at Semester / Year
    c = re.search(
        r"Course Code:\s*([A-Z0-9]+)",
        text,
        re.I
    )
    if c:
        data["course_code"] = c.group(1).strip()

    # Course Name: stop at Year / Subject / end
    n = re.search(
        r"Course Name:\s*(.+?)(?=\s+Year:|\s+Subject|\s+Semester:|$)",
        text,
        re.I
    )
    if n:
        data["course_name"] = n.group(1).strip()

    return data

aqa_meta = extract_aqa_metadata(qb_file)

def extract_co_outcomes(file):
    doc = Document(file)
    co_map = {}

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]

            if len(cells) < 2:
                continue

            co = cells[0].upper()
            desc = cells[1]

            # Detect CO rows purely by pattern
            if re.fullmatch(r"CO\d+", co) and desc:
                co_map[co] = desc

    return co_map

co_outcomes = extract_co_outcomes(qb_file)

def extract_images_from_cell(cell, doc):
    images = []

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            drawing = run._element.xpath(".//w:drawing")
            if not drawing:
                continue

            blips = run._element.xpath(".//a:blip")
            for blip in blips:
                rId = blip.get(qn("r:embed"))
                if rId in doc.part.rels:
                    image_part = doc.part.rels[rId].target_part
                    image_bytes = image_part.blob
                    images.append(BytesIO(image_bytes))

    return images

def set_run_font(run, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold

def write_co_btl_cell(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(8)
    r.bold = False
    r.italic = True

def format_paragraph(p, align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=False):
    p.alignment = align
    for r in p.runs:
        set_run_font(r, bold=bold)

def format_cell(cell, align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=False):
    for p in cell.paragraphs:
        format_paragraph(p, align=align, bold=bold)

def write_cell(cell, text, align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=False, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold

def bold_exact_text(paragraph, target_text):
    """
    Bold ONLY the exact target_text inside a paragraph.
    """
    if target_text not in paragraph.text:
        return

    full_text = paragraph.text
    paragraph.clear()

    before, match, after = full_text.partition(target_text)

    if before:
        r = paragraph.add_run(before)
        set_run_font(r, bold=False)

    r = paragraph.add_run(match)
    set_run_font(r, bold=True)

    if after:
        r = paragraph.add_run(after)
        set_run_font(r, bold=False)

def enforce_tnr_12(cell):
    """
    Force Times New Roman, 12pt on all runs inside a cell
    (equations are automatically ignored by Word)
    """
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = False

# ==================================================
# AQA PARSER 
# ==================================================
def parse_question_bank(file):
    doc = Document(file)
    questions = []
    current_unit = None
    seen_questions = set()  

    def clean(t):
        return re.sub(r"\s+", " ", t).strip()

    def extract_mcq_text(cell):
        lines = []
        opt_count = 0

        for p in cell.paragraphs:
            txt = p.text.strip()
            if txt:
                # Auto-label options if missing
                if opt_count > 0 and not re.match(r"[a-dA-D]\)", txt):
                    txt = f"{chr(96+opt_count)}) {txt}"
                lines.append(txt)
                opt_count += 1

        return "\n".join(lines)

    def is_mcq_cell(cell):
        text = " ".join(p.text for p in cell.paragraphs).lower()
        options = re.findall(r"\b[a-d]\)", text)
        return len(set(options)) >= 3

    for table in doc.tables:
        for row in table.rows:
            cells = [clean(c.text) for c in row.cells]
            if len(cells) < 6:
                continue
            co = cells[1].upper() if re.fullmatch(r"CO\d+", cells[1], re.I) else "CO1"
            k  = cells[2].upper() if re.fullmatch(r"K\d+",  cells[2], re.I) else "K2"

            row_text = " ".join(cells)

            # UNIT
            um = re.search(r"UNIT\s*[-–]?\s*(\d+)", row_text, re.I)
            if um:
                current_unit = f"UNIT-{um.group(1)}"
                continue

            if not current_unit or len(cells) < 6:
                continue

            mark_match = re.search(r"\b(2|4|8|16)\b", cells[4])
            if not mark_match:
                continue

            marks = int(mark_match.group())
            q_cell = row.cells[3]
            images = extract_images_from_cell(q_cell, doc)
            portion = cells[5]

            # PART
            part = (
                "PART A" if marks == 2 else
                "PART B" if marks == 4 else
                "PART B/C" if marks == 8 else
                "PART C" if marks == 16 else None
            )
            if not part:
                continue

            is_mcq = False

            if marks == 2 and is_mcq_cell(q_cell):
                qtext = extract_mcq_text(q_cell)
                is_mcq = True
            else:
    # THEORY / LONG ANSWER → MUST HAVE TEXT
                qtext = q_cell.text.strip()

            if not qtext:
                qtext = " "

            # GLOBAL UNIQUE QUESTION FILTER
            key = (current_unit, marks, qtext if is_mcq else id(q_cell))

            if key in seen_questions:
                continue
            seen_questions.add(key)

            questions.append({
                "uid": f"{current_unit}_{marks}_{len(seen_questions)}",
                "unit": current_unit,
                "part": part,
                "text": qtext,
                "images": images,
                "xml": deepcopy(q_cell._element),
                "marks": marks,
                "portion": portion,
                "is_mcq": is_mcq,
                "co" : co,
                "k" :k
            })

    return questions

bank = parse_question_bank(qb_file)
st.success(f"Loaded {len(bank)} questions")

if not bank:
    st.stop()

units = sorted(set(q["unit"] for q in bank))

def norm(s):
    return re.sub(r"\s+", " ", s.strip())

# ==================================================
# PICKER (NO DUPLICATES)
# ==================================================
def pick(pool, used):
    for q in pool:
        uid = q.get("uid")
        if uid not in used:
            used.add(uid)
            return q
    return None

def delete_row(table, row):
    tbl = table._tbl
    tbl.remove(row._tr)

def find_co_index_table(doc):
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().upper() for c in row.cells]
            if len(cells) >= 2 and cells[0] == "CO INDEX" and "COURSE OUTCOMES" in cells[1]:
                return table
    return None

def has_equation(cell_xml):
    return bool(cell_xml.xpath(".//m:oMath | .//m:oMathPara"))

def pick_pair_same_k(pool_a, pool_b, used):
    """
    Pick one question from pool_a and pool_b
    such that both have the SAME K-level.
    """
    random.shuffle(pool_a)
    random.shuffle(pool_b)

    for qa in pool_a:
        if qa["uid"] in used:
            continue

        for qb in pool_b:
            if qb["uid"] in used:
                continue

            if qa.get("k") == qb.get("k"):
                used.add(qa["uid"])
                used.add(qb["uid"])
                return qa, qb

    return None, None

# ==================================================
# TEMPLATE FILLER 
# ==================================================
def fill_template(template_file, slots, meta, semester, month_duration, qp_code):
    doc = Document(template_file)

    # ---------------- PLACEHOLDER REPLACEMENT ----------------
    replacements = {
        "<department>": meta["department"],
        "<course code>": meta["course_code"],
        "<course name>": meta["course_name"],
        "<sem>": semester,
        "<month duration>": month_duration,
        "<qp_code>" : qp_code
    }

    def replace_placeholders(doc, replacements):

        def replace_in_paragraph(paragraph):

            if not paragraph.runs:
                return

            full_text = "".join(run.text for run in paragraph.runs)
            new_text = full_text

            for key, value in replacements.items():
                if key in new_text:
                    new_text = new_text.replace(key, value)

        # Only rewrite if changed
            if new_text != full_text:
            # Clear all runs safely
                for run in paragraph.runs:
                    run.text = ""

            # Write replaced text in first run
                paragraph.runs[0].text = new_text

    # Normal paragraphs
        for p in doc.paragraphs:
            replace_in_paragraph(p)

    # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p)
    replace_placeholders(doc, replacements)

    # ---------------- GLOBAL FONT & ALIGNMENT ----------------
    for p in doc.paragraphs:
        format_paragraph(p, WD_PARAGRAPH_ALIGNMENT.CENTER)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                format_cell(cell, WD_PARAGRAPH_ALIGNMENT.CENTER)

    # ---------------- QUESTION TABLE ----------------
    qtable = None
    for table in doc.tables:
        text = " ".join(c.text.upper() for r in table.rows for c in r.cells)
        if "PART A" in text and "Q. NO" in text:
            qtable = table
            break

    if not qtable:
        st.error("Question table not found")
        return None

    bloom_totals = {"K1":0,"K2":0,"K3":0,"K4":0}
    used_cos = set()
    rows_to_delete = []

    for row in qtable.rows:
        qkey = None
        for cell in row.cells:
            if re.fullmatch(r"\d+(\s*\([a-z]\))*(\s*\([ivx]+\))*", cell.text.strip(), re.I):
                qkey = cell.text.strip()
                break

        if not qkey:
            continue

        if qkey in slots:
            q = slots[qkey]

            # CO, BTL column (FIRST column only) 
            write_co_btl_cell(
    row.cells[0],
    f"{q.get('co','')}, {q.get('k','')}"
)

            # Question text (SAFE column)
            q_cell = row.cells[-2]
            q_cell._element.clear_content()

            if q.get("xml") is not None and has_equation(q["xml"]):
                for el in q["xml"]:
                    q_cell._element.append(deepcopy(el))

            else:
                q_cell.text = q.get("text", "")

            for img in q.get("images", []):
                p = q_cell.add_paragraph()
                run = p.add_run()
                run.add_picture(img, width=Inches(3))
            enforce_tnr_12(q_cell)

        # Marks 
            write_cell(row.cells[-1], q["marks"], bold=True, size=9)

            # Bloom count
            k = q.get("k","").upper()
            if k in bloom_totals:
                bloom_totals[k] += q["marks"]

            if "co" in q:
                used_cos.add(q["co"])
        else:
            if "(OR)" not in " ".join(c.text.upper() for c in row.cells):
                rows_to_delete.append(row)

    for r in rows_to_delete:
        qtable._tbl.remove(r._tr)

    # ---------------- BLOOM TAXONOMY TABLE ----------------
    for table in doc.tables:

    # Collect row texts
        rows_text = [" ".join(c.text.upper() for c in r.cells) for r in table.rows]

        header_row = None
        value_row = None

    # Step 1: detect header row
        for i, txt in enumerate(rows_text):
            if "REMEMBER" in txt and "UNDERSTAND" in txt and "ANALYZE" in txt:
                header_row = i
                break

        if header_row is None:
            continue

    # Step 2: value row is the NEXT row after header
        if header_row + 1 < len(table.rows):
            value_row = table.rows[header_row + 1]
        else:
            continue

    # Step 3: map columns by header text
        header_cells = table.rows[header_row].cells

        col_map = {}
        for idx, cell in enumerate(header_cells):
            t = cell.text.upper()
            if "K-1" in t:
                col_map["K1"] = idx
            elif "K-2" in t:
                col_map["K2"] = idx
            elif "K-3" in t:
                col_map["K3"] = idx
            elif "K-4" in t:
                col_map["K4"] = idx
            elif "TOTAL" in t:
                col_map["TOTAL"] = idx

    # Step 4: write values safely
        if "K1" in col_map:
            write_cell(value_row.cells[col_map["K1"]], bloom_totals["K1"], bold=True)
        if "K2" in col_map:
            write_cell(value_row.cells[col_map["K2"]], bloom_totals["K2"], bold=True)
        if "K3" in col_map:
            write_cell(value_row.cells[col_map["K3"]], bloom_totals["K3"], bold=True)
        if "K4" in col_map:
            write_cell(value_row.cells[col_map["K4"]], bloom_totals["K4"], bold=True)
        if "TOTAL" in col_map:
            write_cell(value_row.cells[col_map["TOTAL"]], sum(bloom_totals.values()), bold=True)

    # ---------------- CO INDEX TABLE ----------------
    co_table = find_co_index_table(doc)

    if co_table:
        co_list = sorted(used_cos, key=lambda x: int(x[2:]))

        header_idx = None
        bloom_idx = None

    # Locate CO header and Bloom start
        for i, row in enumerate(co_table.rows):
            row_text = " ".join(c.text.upper() for c in row.cells)

            if row.cells[0].text.strip().upper() == "CO INDEX":
                header_idx = i

            if "MARKS DISTRIBUTION" in row_text:
                bloom_idx = i
                break

        if header_idx is not None and bloom_idx is not None:
            co_rows = co_table.rows[header_idx + 1 : bloom_idx]

        # Fill required COs
            for i, co in enumerate(co_list):
                if i >= len(co_rows):
                    break
                write_cell(co_rows[i].cells[0], co, bold=True)
                write_cell(co_rows[i].cells[1], co_outcomes.get(co, ""))

        # DELETE unused CO rows (bottom-up)
            for row in reversed(co_rows[len(co_list):]):
                delete_row(co_table, row)
    if co_table:
        for row in co_table.rows:
        # Skip header rows
            if row.cells[0].text.strip().upper().startswith("CO"):
                p = row.cells[1].paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for r in p.runs:
                    set_run_font(r, bold=False)

    # ---------------- APPLY BOLD PHRASES ----------------
    for p in doc.paragraphs:
        for phrase in BOLD_PHRASES:
            if phrase.upper() in p.text.upper():
                for r in p.runs:
                    set_run_font(r, bold=True)

    for table in doc.tables:
        for row in table.rows:
            for idx, cell in enumerate(row.cells):

            # Skip CO, BTL column
                if idx == 0:
                    continue

                for p in cell.paragraphs:
                    for phrase in BOLD_PHRASES:
                        if phrase.upper() in p.text.upper():
                            for r in p.runs:
                                set_run_font(r, bold=True)

    if qtable:
        for row in qtable.rows:
            cell = row.cells[0]
            txt = cell.text.strip().upper()

        # Match only patterns like "CO1, K3"
            if re.fullmatch(r"CO\d+\s*,\s*K\d+", txt):
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                r = p.add_run(txt)
                r.font.name = "Times New Roman"
                r.font.size = Pt(8)
                r.bold = False
                r.italic = True

    # ---------------- FINAL BOLD: COURSE LINE & MONTH ----------------
    course_line = f"{meta['course_code']} - {meta['course_name']}"

    for p in doc.paragraphs:
        bold_exact_text(p, course_line)
        bold_exact_text(p, month_duration)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    bold_exact_text(p, course_line)
                    bold_exact_text(p, month_duration)

    # ---------------- HEADER FIXES ----------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip().upper()

            # CO INDEX → BOLD
                if txt == "CO INDEX":
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for r in p.runs:
                            r.font.bold = True
                            r.font.size = Pt(12)

            # COURSE OUTCOMES → CENTER
                if txt == "COURSE OUTCOMES":
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # CO, BTL → CENTER
                if txt == "CO, BTL":
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ==================================================
# UI INPUTS
# ==================================================

# ---------- PART A ----------
st.header("PART A – 10 Questions (2 Marks)")
st.info("Q1–10 are auto-generated based on fixed UNIT mapping.")

# ---------- PART B ----------
st.header("PART B – 5 x 16 = 80 Marks")

pb_config = []

units_sorted = sorted(units, key=lambda x: int(x.split("-")[1]))

for i in range(5):

    qno = 11 + i
    unit = units_sorted[i]

    st.subheader(f"Question {qno} → {unit}")

    qtype = st.selectbox(
        f"Q{qno} Type",
        ["Single", "Split"],
        key=f"type_{qno}"
    )

    portion_a = st.selectbox(
        f"Q{qno} (a) Portion",
        ["I", "II"],
        key=f"pa_{qno}"
    )

    portion_b = st.selectbox(
        f"Q{qno} (b) Portion",
        ["I", "II"],
        key=f"pb_{qno}"
    )

    pb_config.append((unit, qtype, portion_a, portion_b))
    
# ==================================================
# GENERATE
# ==================================================
if st.button("Generate Question Paper"):

    used = set()
    slots = {}

    # ==================================================
    # PART A – FIXED UNIT MAPPING
    # ==================================================

    unit_map = {
    1: "UNIT-1", 2: "UNIT-1",
    3: "UNIT-2", 4: "UNIT-2",
    5: "UNIT-3", 6: "UNIT-3",
    7: "UNIT-4", 8: "UNIT-4",
    9: "UNIT-5", 10: "UNIT-5"
    }

    random.shuffle(bank)

    for qno in range(1, 11):

        unit = unit_map[qno]

        pool = [
        q for q in bank
        if q["unit"] == unit and q["marks"] == 2
        ]

        if qno % 2 == 0:
            pool = [q for q in pool if q["is_mcq"]]
        else:
            pool = [q for q in pool if not q["is_mcq"]]

        random.shuffle(pool)
        
        q = pick(pool, used)

        slots[str(qno)] = q if q else {
        "text": "Not available",
        "marks": 2
        }

    # ==================================================
    # PART B – AUTO SAME K MATCHING
    # ==================================================

    qno = 11

    for unit, qtype, portion_a, portion_b in pb_config:
    
        if qtype == "Split":

            pool_a = [
            q for q in bank
            if q["unit"] == unit
            and q["marks"] == 8
            and q["portion"] == portion_a
            ]

            pool_b = [
            q for q in bank
            if q["unit"] == unit
            and q["marks"] == 8
            and q["portion"] == portion_b
            ]

        # First pair (i)
            a1, b1 = pick_pair_same_k(pool_a, pool_b, used)

        # Second pair (ii)
            a2, b2 = pick_pair_same_k(pool_a, pool_b, used)

            slots[f"{qno} (a) (i)"] = a1 if a1 else {"text": "Not available", "marks": 8}
            slots[f"{qno} (b) (i)"] = b1 if b1 else {"text": "Not available", "marks": 8}
            slots[f"{qno} (a) (ii)"] = a2 if a2 else {"text": "Not available", "marks": 8}
            slots[f"{qno} (b) (ii)"] = b2 if b2 else {"text": "Not available", "marks": 8}

        else:  # Single 16

            pool_a = [
            q for q in bank
            if q["unit"] == unit
            and q["marks"] == 16
            and q["portion"] == portion_a
            ]

            pool_b = [
            q for q in bank
            if q["unit"] == unit
            and q["marks"] == 16
            and q["portion"] == portion_b
            ]

            qa, qb = pick_pair_same_k(pool_a, pool_b, used)

            slots[f"{qno} (a)"] = qa if qa else {"text": "Not available", "marks": 16}
            slots[f"{qno} (b)"] = qb if qb else {"text": "Not available", "marks": 16}

        qno += 1

    doc = fill_template(
        template_file,
        slots,
        meta=aqa_meta,
        semester=semester,
        month_duration=month_duration,
        qp_code=qp_code
    )

    if doc:
        st.download_button(
            "Download Question Paper",
            doc.getvalue(),
            "KCET_Semester_QP.docx"
        )
        st.success("Semester Question Paper Generated Successfully")